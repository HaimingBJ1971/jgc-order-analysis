"""
Word report generator for period comparison analysis.
Mirrors the PDF content structure.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xC0, 0x39, 0x2B)
GREEN = RGBColor(0x27, 0xAE, 0x60)
BLACK = RGBColor(0x00, 0x00, 0x00)


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '等线'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')


def _add_para(doc, text, size=9, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '等线'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    return p


def _add_table(doc, headers, rows, col_widths=None):
    """Add a styled table. headers: list of str, rows: list of lists of (text, color_or_None, align)."""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    # Mark header row to repeat across pages
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement('w:tblHeader'))
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        run.font.name = '等线'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Dark background
        shading = cell._element.get_or_add_tcPr()
        s = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '2c3e50', qn('w:val'): 'clear'
        })
        shading.append(s)
        run.font.color.rgb = RGBColor(255, 255, 255)
    # Data rows
    for ri, row in enumerate(rows):
        for ci, cell_data in enumerate(row):
            text = cell_data[0] if isinstance(cell_data, tuple) else str(cell_data)
            color = cell_data[1] if isinstance(cell_data, tuple) and len(cell_data) > 1 else None
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(8)
            run.font.name = '等线'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
            if color:
                run.font.color.rgb = color
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
    return table


def _change_str(val, is_neg):
    """Format a change value with color. Returns (text, color)."""
    if is_neg:
        return (val, RED)
    else:
        return (val, GREEN)


def _qty_change(cur, base):
    """Return ((diff_text, color), (pct_text, color)) for quantity comparisons."""
    if base and base > 0:
        diff = cur - base
        rate = round(diff / base * 100, 1)
        return (
            _change_str(f'{diff:+d}', diff < 0),
            _change_str(f'{rate:+.1f}%', diff < 0),
        )
    return ('-', None), ('-', None)

def generate_comparison_word(output_path, period_info, comparison_info, comp_data, mode, store_name):
    doc = Document()

    # ── Title ──
    _add_heading(doc, '周期对比分析报告', level=0)
    _add_para(doc, f'{store_name} ｜ 本期：{period_info["period_label"]}', size=12, bold=True)

    ringbi_label = comparison_info.get('ringbi_label', '环比')
    tongbi_label = comparison_info.get('tongbi_label', '同比')
    ringbi_dates = f'{comparison_info["ringbi_start"]} ~ {comparison_info["ringbi_end"]}'
    tongbi_dates = f'{comparison_info["tongbi_start"]} ~ {comparison_info["tongbi_end"]}'
    _add_para(doc, f'环比：{ringbi_label}（{ringbi_dates}）  ｜  同比：{tongbi_label}（{tongbi_dates}）', size=9)

    # ── Section 1 ──
    _add_heading(doc, '一、经营数据对比', level=1)
    headers = ['指标', '本期', '环比', '', '同比', '']
    sub_headers = ['', '', '变化', '变化%', '变化', '变化%']

    rows = [sub_headers]
    for item in comp_data['operational']:
        ring_neg = item['ringbi_diff'].startswith('-')
        tong_neg = item['tongbi_diff'].startswith('-')
        rows.append([
            (item['label'], None),
            (item['current'], None),
            _change_str(item['ringbi_diff'], ring_neg),
            _change_str(item['ringbi_pct'], ring_neg),
            _change_str(item['tongbi_diff'], tong_neg),
            _change_str(item['tongbi_pct'], tong_neg),
        ])
    _add_table(doc, headers, rows)
    _add_para(
        doc,
        '注：整体营业额 = 堂食分桌总营业额 + 自取外卖单、吧台及零食购买团体、第三方平台外卖单合计；第三方平台外卖按已完成订单的订单收入计入。',
        size=7,
        color=RGBColor(0x66, 0x66, 0x66),
    )

    # ── Section 2: Drinks ──
    _add_heading(doc, '二、酒水饮料甜品销售排行', level=1)
    drinks_cur = comp_data.get('drinks_current', [])
    drinks_ring = {name: (info['qty'] if isinstance(info, dict) else info)
                   for name, info in comp_data.get('drinks_ringbi', [])}
    drinks_tong = {name: (info['qty'] if isinstance(info, dict) else info)
                   for name, info in comp_data.get('drinks_tongbi', [])}
    if drinks_cur:
        d_headers = ['商品分类', '商品名称', '本期销量', '环比销量', '变化', '变化率', '同比销量', '变化', '变化率']
        d_rows = []
        for name, info in drinks_cur:
            qty = info['qty'] if isinstance(info, dict) else info
            cat_name = info.get('cat', '') if isinstance(info, dict) else ''
            ring_qty = drinks_ring.get(name)
            tong_qty = drinks_tong.get(name)
            rd, rr = _qty_change(qty, ring_qty)
            td, tr = _qty_change(qty, tong_qty)
            d_rows.append([
                (cat_name, None), (name, None), (str(qty), None),
                (str(ring_qty) if ring_qty else '-', None), rd, rr,
                (str(tong_qty) if tong_qty else '-', None), td, tr,
            ])
        _add_table(doc, d_headers, d_rows)
        _add_para(
            doc,
            '注：本表销量采用 POS 店内全量正收入口径，包含自取外卖单、吧台及零食购买团体；赠送、免单、全额优惠等菜品收入≤0的商品不计入销量。不得用“销售数量-赠菜数量”替代正收入销量。第三方平台外卖若无商品级明细，不纳入商品销量排行。',
            size=7,
            color=RGBColor(0x66, 0x66, 0x66),
        )
    else:
        _add_para(doc, '（无数据）')

    # ── Section 3: Dishes ──
    _add_heading(doc, '三、重点菜品对比', level=1)
    cur_dishes = comp_data['dishes_current']
    ring_d_dict = dict(comp_data['dishes_ringbi'])
    tong_d_dict = dict(comp_data['dishes_tongbi'])
    dish_headers = ['菜品名称', '本期销量', '环比销量', '变化', '变化率', '同比销量', '变化', '变化率']
    dish_rows = []
    for name, qty in cur_dishes:
        rq = ring_d_dict.get(name); tq = tong_d_dict.get(name)
        rd, rr = _qty_change(qty, rq); td, tr = _qty_change(qty, tq)
        dish_rows.append([
            (name, None), (str(qty), None),
            (str(rq) if rq else '-', None), rd, rr,
            (str(tq) if tq else '-', None), td, tr,
        ])
    _add_table(doc, dish_headers, dish_rows)
    _add_para(
        doc,
        '注：本表销量采用 POS 店内全量正收入口径，包含自取外卖单、吧台及零食购买团体；赠送、免单、全额优惠等菜品收入≤0的商品不计入销量。不得用“销售数量-赠菜数量”替代正收入销量。第三方平台外卖若无商品级明细，不纳入商品销量排行。',
        size=7,
        color=RGBColor(0x66, 0x66, 0x66),
    )

    # ── Section 4: Categories ──
    cat_dim = comp_data.get('category_dimension') or {}
    section_title = cat_dim.get('title', '四、商品中类销售额分布')
    col_label = cat_dim.get('column', '商品中类')
    cat_field = cat_dim.get('field', '商品中类')
    _add_heading(doc, section_title, level=1)
    cats_cur_list = comp_data.get('cats_current', [])
    cats_ring = dict(comp_data.get('cats_ringbi', []))
    cats_tong = dict(comp_data.get('cats_tongbi', []))
    if cats_cur_list:
        cur_total = sum(v for _, v in cats_cur_list)
        # ring_total computed on the fly for rate
        cat_headers = [col_label, '本期金额', '本期占比', '环比变化', '环比变化率', '同比变化', '同比变化率']
        cat_rows = []
        for cat_name, rev in cats_cur_list:
            ring_rev = cats_ring.get(cat_name)
            tong_rev = cats_tong.get(cat_name)
            cur_pct = round(rev / cur_total * 100, 1) if cur_total > 0 else 0
            if ring_rev and ring_rev > 0:
                rd_val = rev - ring_rev; rr_val = round(rd_val / ring_rev * 100, 1)
                rd_s = _change_str(f'{rd_val:+,.0f}', rd_val < 0)
                rr_s = _change_str(f'{rr_val:+.1f}%', rd_val < 0)
            else:
                rd_s = ('-', None); rr_s = ('-', None)
            if tong_rev and tong_rev > 0:
                td_val = rev - tong_rev; tr_val = round(td_val / tong_rev * 100, 1)
                td_s = _change_str(f'{td_val:+,.0f}', td_val < 0)
                tr_s = _change_str(f'{tr_val:+.1f}%', td_val < 0)
            else:
                td_s = ('-', None); tr_s = ('-', None)
            cat_rows.append([
                (cat_name, None), (f'¥{rev:,.0f}', None), (f'{cur_pct}%', None),
                rd_s, rr_s, td_s, tr_s,
            ])
        # Totals row
        cur_total_rev = sum(v for _, v in cats_cur_list)
        ring_total_rev = sum(v for _, v in comp_data.get('cats_ringbi', []))
        tong_total_rev = sum(v for _, v in comp_data.get('cats_tongbi', []))
        if ring_total_rev > 0:
            rtd = cur_total_rev - ring_total_rev
            cat_rows.append([
                ('合计', None), (f'¥{cur_total_rev:,.0f}', None), ('100%', None),
                _change_str(f'{rtd:+,.0f}', rtd < 0),
                _change_str(f'{rtd / ring_total_rev * 100:+.1f}%', rtd < 0),
                _change_str(f'{cur_total_rev - tong_total_rev:+,.0f}', (cur_total_rev - tong_total_rev) < 0) if tong_total_rev > 0 else ('-', None),
                _change_str(f'{(cur_total_rev - tong_total_rev) / tong_total_rev * 100:+.1f}%', (cur_total_rev - tong_total_rev) < 0) if tong_total_rev > 0 else ('-', None),
            ])
        _add_table(doc, cat_headers, cat_rows)
        _add_para(
            doc,
            '注：以上为 POS 店内商品归因销售额，包含自取外卖单、吧台及零食购买团体；赠送、免单、全额优惠等菜品收入≤0的商品不计入。套餐父项不计入，套餐子项按实际商品品类归因，避免父项和子项重复统计。第三方平台外卖若无商品级明细，不纳入商品分类销售额。',
            size=7,
            color=RGBColor(0x66, 0x66, 0x66),
        )
        uncat = comp_data.get('uncategorized_products') or []
        if uncat:
            from comparator import format_uncategorized_note
            for line in format_uncategorized_note(uncat, cat_field).split('\n'):
                _add_para(doc, line, size=6, color=RGBColor(0x66, 0x66, 0x66))
    else:
        _add_para(doc, '（无数据）')

    # ── Section 5: Buckets ──
    _add_heading(doc, '五、客单价区间对比', level=1)
    buckets_cur = comp_data['buckets_current']
    buckets_ring = comp_data['buckets_ringbi']
    buckets_tong = comp_data['buckets_tongbi']
    bucket_order = ['≥300', '200~300', '150~200', '100~150', '<100']
    bkt_headers = ['客单价区间', '本期单数', '本期占比', '环比单数变化', '环比变化率', '同比单数变化', '同比变化率']
    bkt_rows = []
    for bk in bucket_order:
        cur_cnt = buckets_cur.get(bk, {}).get('订单数', 0)
        cur_pct = buckets_cur.get(bk, {}).get('占比', 0)
        ring_cnt = buckets_ring.get(bk, {}).get('订单数') if buckets_ring else None
        tong_cnt = buckets_tong.get(bk, {}).get('订单数') if buckets_tong else None
        if ring_cnt and ring_cnt > 0:
            rd_val = cur_cnt - ring_cnt; rr_val = round(rd_val / ring_cnt * 100, 1)
            rd_s = _change_str(f'{rd_val:+d}', rd_val < 0)
            rr_s = _change_str(f'{rr_val:+.1f}%', rd_val < 0)
        else:
            rd_s = ('-', None); rr_s = ('-', None)
        if tong_cnt and tong_cnt > 0:
            td_val = cur_cnt - tong_cnt; tr_val = round(td_val / tong_cnt * 100, 1)
            td_s = _change_str(f'{td_val:+d}', td_val < 0)
            tr_s = _change_str(f'{tr_val:+.1f}%', td_val < 0)
        else:
            td_s = ('-', None); tr_s = ('-', None)
        bkt_rows.append([
            (bk, None), (str(cur_cnt), None), (f'{cur_pct}%', None),
            rd_s, rr_s, td_s, tr_s,
        ])
    _add_table(doc, bkt_headers, bkt_rows)
    _add_para(
        doc,
        '注：本部分客单价对应“堂食分桌总营业额”，即整体营业额扣除自取外卖单、吧台及零食购买团体、第三方平台外卖单合计。',
        size=7,
        color=RGBColor(0x66, 0x66, 0x66),
    )

    # ── Data quality note ──
    dq = comp_data['data_quality']
    warnings = []
    if dq['ringbi_missing']:
        warnings.append(f'环比数据缺失（{comparison_info["ringbi_label"]}在数据库中无记录）')
    if dq['tongbi_missing']:
        warnings.append(f'同比数据缺失（{comparison_info["tongbi_label"]}在数据库中无记录）')
    if warnings:
        _add_heading(doc, '数据完整性说明', level=1)
        for w in warnings:
            _add_para(doc, w, size=9, color=RED)

    doc.save(output_path)
    print(f"Word 报告已生成: {output_path}")
